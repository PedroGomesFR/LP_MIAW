import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Inches

class App:
    def __init__(self, master):
        self.master = master
        master.title("Insert Images into Word")

        self.label = tk.Label(master, text="Select a folder with images to insert into Word:")
        self.label.pack(pady=10)

        self.select_button = tk.Button(master, text="Select Image Folder", command=self.select_folder)
        self.select_button.pack(pady=5)

        self.output_button = tk.Button(master, text="Create Word Document with Images", command=self.create_document)
        self.output_button.pack(pady=5)

        self.status_label = tk.Label(master, text="")
        self.status_label.pack(pady=20)

        self.selected_folder = ""

    def select_folder(self):
        """Sélectionne un dossier contenant les images."""
        self.selected_folder = filedialog.askdirectory()
        if self.selected_folder:
            self.status_label.config(text=f"Selected folder: {os.path.basename(self.selected_folder)}")

    def create_document(self):
        """Crée un document Word avec les images sélectionnées."""
        if not self.selected_folder:
            messagebox.showwarning("Warning", "Please select a folder first!")
            return

        # Demander où sauvegarder le fichier Word généré
        output_docx_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                        filetypes=[("Word files", "*.docx")])
        if not output_docx_path:
            return

        try:
            # Créer un document Word
            doc = Document()
            doc.add_heading('Images Inserted into Document', level=1)

            # Créer un tableau pour insérer les images et leurs noms
            columns = 2  # Deux colonnes: une pour l'image, une pour le nom
            table = doc.add_table(rows=1, cols=columns)  # Une ligne pour l'entête
            table.style = 'Table Grid'

            # Ajouter un en-tête (optionnel)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Image'
            hdr_cells[1].text = 'Image Name'

            # Trier les fichiers pour les ajouter dans l'ordre
            image_files = [f for f in os.listdir(self.selected_folder) if f.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp'))]
            image_files.sort()

            # Ajouter les images et leurs noms au tableau
            row_cells = None
            for idx, image_file in enumerate(image_files):
                # Nouvelle ligne après chaque image
                row_cells = table.add_row().cells

                # Ajouter l'image dans la première colonne
                paragraph = row_cells[0].paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(os.path.join(self.selected_folder, image_file), width=Inches(0.6), height=Inches(0.6))  # Taille des images

                # Ajouter le nom de l'image dans la deuxième colonne
                row_cells[1].text = image_file

            # Sauvegarder le fichier Word avec les images et leurs noms
            doc.save(output_docx_path)
            
            self.status_label.config(text="Document created successfully!")
            messagebox.showinfo("Success", "Document created successfully!")

        except Exception as e:
            messagebox.showerror("Error", f"An error occurred: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = App(root)
    root.mainloop()
